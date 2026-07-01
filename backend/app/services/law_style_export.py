"""
律所级合同版式导出服务 - 专业版
==================================
- 清洁版 (clean): 应用 AI 建议后重排，无痕迹
- 原始版 (original): 原文件解析后重排律所级版式
- 修改版 (modified): 原合同正文 + Word 批注框 (w:comment) + 末尾修改对照表

律所级版式要素:
1. 封面页 (律所LOGO位 / 合同名 / 编号 / 日期 / 起草律师)
2. 目录 (基于标题解析)
3. 修改对照表 (修改版专用)
4. 页眉: 左律所名 + 中合同名 + 右第 X 页 / 共 Y 页
5. 页脚: "本文件由XX律所出具  编号XXX"
6. 装订线 (左页边距灰色细竖线)
7. 正文: 第X条 / X.Y / X.Y.Z 律所通行编号
8. 签字盖章: 甲方乙方分两列
9. 附件区

字号:
- 合同标题: 黑体 22pt 居中
- 第X条: 黑体 16pt
- X.Y (二级): 黑体 14pt
- X.Y.Z (三级): 黑体 12pt 加粗
- 正文: 宋体 12pt 1.5 倍行距 2 字符缩进
- 页眉/页脚: 宋体 9pt
"""
import io
import os
import re
from typing import List, Dict, Optional, Tuple


# ============== 工具函数 ==============

CN_NUM = {1:'一',2:'二',3:'三',4:'四',5:'五',6:'六',7:'七',8:'八',9:'九',10:'十',11:'十一',12:'十二',13:'十三',14:'十四',15:'十五',16:'十六',17:'十七',18:'十八',19:'十九',20:'二十'}

# 文档信息默认值
DEFAULT_LAWYER = "经办律师"
DEFAULT_DATE = "2026年6月17日"
CONTRACT_NO_PREFIX = "CR"

def _cn_num(n: int) -> str:
    if n <= 20:
        return CN_NUM.get(n, str(n))
    if n < 100:
        a, b = divmod(n, 10)
        return f"{CN_NUM[a]}十{CN_NUM.get(b, str(b)) if b else ''}"
    return str(n)


def _risk_label(level: str) -> str:
    return {"high": "高风险", "medium": "中等风险", "low": "低风险"}.get(level, "未分级")


def _risk_color(level: str) -> Tuple[int, int, int]:
    return {"high": (0xC0, 0x00, 0x00), "medium": (0xC8, 0x80, 0x00), "low": (0x00, 0x80, 0x00)}.get(level, (0x00, 0x00, 0x00))


def _detect_clause_level(text: str) -> Tuple[int, str]:
    """判断条款层级: 0=合同标题 1=第X条 2=一、 3=(一) 4=1. 5=(1) 9=正文
    也支持 Markdown 标题: # → 0, ## → 1, ### → 2, #### → 3
    """
    t = text.strip()
    # Markdown 标题 (优先检测, 不去掉 # 前缀)
    if re.match(r'^#{1}\s+', t):
        return 0, re.sub(r'^#\s+', '', t)
    if re.match(r'^#{2}\s+', t):
        return 1, re.sub(r'^#{2}\s+', '', t)
    if re.match(r'^#{3}\s+', t):
        return 2, re.sub(r'^#{3}\s+', '', t)
    if re.match(r'^#{4,6}\s+', t):
        return 3, re.sub(r'^#{4,6}\s+', '', t)
    # 中文合同条款
    if re.match(r'^第[一二三四五六七八九十百零〇\d]+条', t):
        return 1, t
    if re.match(r'^[一二三四五六七八九十]+、', t):
        return 2, t
    if re.match(r'^（[一二三四五六七八九十]+）', t):
        return 3, t
    if re.match(r'^\d+\.\s*\S', t) or re.match(r'^\d+\.', t):
        return 4, t
    if re.match(r'^（\d+）', t) or re.match(r'^\(\d+\)', t):
        return 5, t
    return 9, t


# ============== DOCX 基础工具 ==============

def _setup_docx_default_style(doc):
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    s = doc.sections[0]
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18)
    s.right_margin = Cm(2.54)
    s.page_height = Cm(29.7)
    s.page_width = Cm(21.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '宋体')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = 3
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _set_run_font(run, eastasia='宋体', size=12, bold=False, color=None):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), eastasia)
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')


def _add_blank(doc, count=1, size=12):
    for _ in range(count):
        p = doc.add_paragraph()
        run = p.add_run('')
        _set_run_font(run, '宋体', size)


def _add_para(doc, text, level=9, bold=False, align='left', first_line_indent=True, space_before=0, color=None):
    from docx.shared import Cm, Pt, RGBColor
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = 1
    elif align == 'right':
        p.alignment = 2
    elif align == 'justify':
        p.alignment = 3
    pf = p.paragraph_format
    if first_line_indent and level == 9:
        pf.first_line_indent = Cm(0.74)
    if space_before:
        pf.space_before = Pt(space_before)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = 3
    
    # 确定字体和大小
    if level == 0:
        fname, fsize, fbold = '黑体', 22, True
    elif level == 1:
        fname, fsize, fbold = '黑体', 16, True
        pf.space_before = Pt(12); pf.first_line_indent = Cm(0)
    elif level == 2:
        fname, fsize, fbold = '黑体', 14, True
        pf.space_before = Pt(6); pf.first_line_indent = Cm(0)
    elif level == 3:
        fname, fsize, fbold = '黑体', 12, True
    elif level == 4:
        fname, fsize, fbold = '宋体', 12, True
    elif level == 5:
        fname, fsize, fbold = '宋体', 12, False
    else:
        fname, fsize, fbold = '宋体', 12, bold
    
    # 解析 **加粗** 和 ~~删除线~~ 标记
    # 加粗 → 蓝色, 删除线 → 红色+删除线
    segments = []
    remaining = text
    while remaining:
        # 先找 **加粗**
        bold_match = re.search(r'\*\*(.+?)\*\*', remaining)
        # 再找 ~~删除线~~
        del_match = re.search(r'~~(.+?)~~', remaining)
        
        # 取最先出现的
        bold_pos = bold_match.start() if bold_match else float('inf')
        del_pos = del_match.start() if del_match else float('inf')
        
        if bold_pos == float('inf') and del_pos == float('inf'):
            segments.append(('normal', remaining))
            break
        elif bold_pos <= del_pos:
            if bold_pos > 0:
                segments.append(('normal', remaining[:bold_pos]))
            segments.append(('bold', bold_match.group(1)))
            remaining = remaining[bold_match.end():]
        else:
            if del_pos > 0:
                segments.append(('normal', remaining[:del_pos]))
            segments.append(('del', del_match.group(1)))
            remaining = remaining[del_match.end():]
    
    # 渲染 segments
    for seg_type, seg_text in segments:
        run = p.add_run(seg_text)
        if seg_type == 'bold':
            _set_run_font(run, fname, fsize, True, RGBColor(0x00, 0x66, 0xCC))
        elif seg_type == 'del':
            _set_run_font(run, fname, fsize, False, RGBColor(0xCC, 0x00, 0x00))
            run.font.strike = True
        else:
            _set_run_font(run, fname, fsize, fbold, color)
    
    return p


# ============== 页眉/页脚/装订线 ==============

def _insert_field(paragraph, field_code, font_name='宋体', font_size=9):
    """插入 Word 字段 (PAGE / NUMPAGES / DATE 等)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = paragraph._p
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_code} '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    rfonts1 = OxmlElement('w:rFonts')
    rfonts1.set(qn('w:eastAsia'), font_name)
    rPr1.append(rfonts1)
    sz1 = OxmlElement('w:sz')
    sz1.set(qn('w:val'), str(font_size * 2))
    rPr1.append(sz1)
    run1.append(rPr1)
    run1.append(fldChar1)
    run1.append(instrText)
    run1.append(fldChar2)
    run2 = OxmlElement('w:r')
    rPr2 = OxmlElement('w:rPr')
    rfonts2 = OxmlElement('w:rFonts')
    rfonts2.set(qn('w:eastAsia'), font_name)
    rPr2.append(rfonts2)
    sz2 = OxmlElement('w:sz')
    sz2.set(qn('w:val'), str(font_size * 2))
    rPr2.append(sz2)
    run2.append(rPr2)
    t = OxmlElement('w:t')
    t.text = '1'
    run2.append(t)
    run3 = OxmlElement('w:r')
    run3.append(fldChar3)
    p.append(run1)
    p.append(run2)
    p.append(run3)


def _setup_header_footer(doc, contract_title: str, contract_no: str):
    """律所版页眉 + 页脚 + 装订线"""
    from docx.shared import Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.different_first_page_header_footer = True  # 封面无页眉
    
    # 默认页眉 (非首页)
    header = section.header
    h_p = header.paragraphs[0]
    pf = h_p.paragraph_format
    pf.tab_stops.add_tab_stop(Cm(7.5), 1)
    pf.tab_stops.add_tab_stop(Cm(15.5), 2)
    # (律所名已移除)
    run1 = h_p.add_run("\t"); _set_run_font(run1, '宋体', 9)
    run2 = h_p.add_run(contract_title); _set_run_font(run2, '宋体', 9)
    run2.add_tab()
    run3 = h_p.add_run("第 "); _set_run_font(run3, '宋体', 9)
    _insert_field(h_p, 'PAGE', '宋体', 9)
    run4 = h_p.add_run(" 页 / 共 "); _set_run_font(run4, '宋体', 9)
    _insert_field(h_p, 'NUMPAGES', '宋体', 9)
    run5 = h_p.add_run(" 页"); _set_run_font(run5, '宋体', 9)
    # 底部下划线
    pPr = h_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # 页脚
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = 1
    foot_run = f_p.add_run(f"编号: {contract_no}  |  合同审查报告")
    _set_run_font(foot_run, '宋体', 8)
    foot_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    
    # 装订线: page border left
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn('w:pgBorders'))
    if pgBorders is not None:
        sectPr.remove(pgBorders)
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '12')
    left_border.set(qn('w:space'), '24')
    left_border.set(qn('w:color'), 'BFBFBF')
    pgBorders.append(left_border)
    sectPr.append(pgBorders)


# ============== 封面页 ==============

def _add_cover_page(doc, contract_title: str, contract_no: str, lawyer: str, date_str: str, risk_level: str = ""):
    """合同审查报告封面 (通用商务版式)"""
    from docx.shared import RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # 顶部品牌色装饰条
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '36')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _add_blank(doc, 3)
    # 文档类型
    p = doc.add_paragraph(); p.alignment = 1
    run = p.add_run("合同审查报告")
    _set_run_font(run, '黑体', 36, True)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _add_blank(doc, 1)
    if risk_level:
        p = doc.add_paragraph(); p.alignment = 1
        run = p.add_run(f"风险等级: {_risk_label(risk_level)}")
        _set_run_font(run, '黑体', 16, True)
        run.font.color.rgb = RGBColor(*_risk_color(risk_level))
    _add_blank(doc, 2)
    # 合同标题区
    p = doc.add_paragraph(); p.alignment = 1
    run = p.add_run("— 关于 —")
    _set_run_font(run, '宋体', 12)
    p = doc.add_paragraph(); p.alignment = 1
    run = p.add_run(f"《{contract_title}》")
    _set_run_font(run, '黑体', 22, True)
    _add_blank(doc, 4)
    # 元信息表 (带边框)
    table = doc.add_table(rows=3, cols=2)
    table.alignment = 1
    table.autofit = False
    for i, (k, v) in enumerate([
        ("文件编号", contract_no),
        ("审查日期", date_str),
        ("审查人员", lawyer),
    ]):
        c0 = table.rows[i].cells[0]; c0.paragraphs[0].text = ""
        c1 = table.rows[i].cells[1]; c1.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(k); _set_run_font(r0, '黑体', 11, True)
        r1 = c1.paragraphs[0].add_run(v); _set_run_font(r1, '宋体', 11)
        c0.width = Cm(3.5); c1.width = Cm(8.5)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'BFBFBF')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            tcPr.append(shd)
    _add_blank(doc, 6)
    # 底部装饰线 + 副标
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '18')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '1F3864')
    pBdr.append(top)
    pPr.append(pBdr)
    p = doc.add_paragraph(); p.alignment = 1
    run = p.add_run("Contract Review Report")
    _set_run_font(run, '宋体', 10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


# ============== 目录页 ==============

def _add_toc_page(doc, sections: List[str]):
    """目录页 - 引导线 + 右侧页码占位"""
    p = doc.add_paragraph(); p.alignment = 1
    run = p.add_run("目  录"); _set_run_font(run, '黑体', 22, True)
    _add_blank(doc, 1)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _add_blank(doc, 1)
    for sec in sections:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.line_spacing_rule = 3
        p.paragraph_format.space_after = __import__('docx').shared.Pt(2)
        run = p.add_run(sec)
        _set_run_font(run, '宋体', 12)
    doc.add_page_break()


# ============== 正文渲染 ==============

def _is_md_table_row(s: str) -> bool:
    """判断是否为 Markdown 表格行"""
    return s.startswith('|') and s.endswith('|') and s.count('|') >= 2


def _is_md_table_separator(s: str) -> bool:
    """判断是否为 Markdown 表格分隔行 | --- | --- |"""
    if not _is_md_table_row(s):
        return False
    cells = s.strip('|').split('|')
    for cell in cells:
        cell = cell.strip()
        if not cell:
            return False
        if set(cell) - set('-: '):
            return False
    return True


def _parse_md_table_row(s: str) -> List[str]:
    """解析 Markdown 表格行为单元格列表"""
    # 去掉首尾 |, 按 | 分割
    cells = s.strip('|').split('|')
    return [c.strip() for c in cells]


def _add_md_table_to_docx(doc, rows: List[List[str]]):
    """将 Markdown 表格行列表转为 Word 表格"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if len(rows) < 2:
        return
    num_cols = len(rows[0])
    # 确保所有行列数一致
    for r in rows:
        while len(r) < num_cols:
            r.append('')
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row[:num_cols]):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].text = ''
            # 清理 markdown 残留标记
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            run = cell.paragraphs[0].add_run(clean)
            if i == 0:
                _set_run_font(run, '黑体', 10.5, True)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _set_run_font(run, '宋体', 10.5)
    _add_blank(doc, 1)


def _render_body(doc, text: str):
    """渲染合同正文 - 智能识别层级 (支持 Markdown 标题和中文条款编号)"""
    lines = text.split('\n')
    current_level = 9
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            _add_blank(doc, 1)
            i += 1
            continue
        # 检测 Markdown 表格
        if _is_md_table_row(s):
            table_lines = []
            while i < len(lines) and _is_md_table_row(lines[i].strip()):
                row_s = lines[i].strip()
                if not _is_md_table_separator(row_s):
                    table_lines.append(_parse_md_table_row(row_s))
                i += 1
            if len(table_lines) >= 2:
                _add_md_table_to_docx(doc, table_lines)
            else:
                # 不到2行, 当普通文本
                for r in table_lines:
                    _add_para(doc, ' | '.join(r), level=9)
            continue
        # 只去掉 markdown 加粗/代码标记, 保留 # 标题前缀给 _detect_clause_level 识别
        s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
        s = re.sub(r'`([^`]+)`', r'\1', s)
        level, cleaned = _detect_clause_level(s)
        _add_para(doc, cleaned, level=level)
        i += 1


# ============== 签字盖章 ==============

def _add_signature_block(doc, party_a: str = "（以下简称“甲方”）", party_b: str = "（以下简称“乙方”）"):
    """签字盖章块 (通用商务版式)"""
    from docx.shared import Cm
    doc.add_page_break()
    _add_para(doc, "签 字 盖 章", level=0, align='center', first_line_indent=False)
    _add_blank(doc, 1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    # 标题行
    table.rows[0].cells[0].text = ''
    table.rows[0].cells[1].text = ''
    for i, (label_a, label_b) in enumerate([
        ("甲  方:", "乙  方:"),
        ("单位名称:", "单位名称:"),
        ("法定代表人:", "法定代表人:"),
        ("委托代理人:", "委托代理人:"),
        ("日  期:    年   月   日", "日  期:    年   月   日"),
    ]):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.paragraphs[0].text = ''
        c1.paragraphs[0].text = ''
        r0 = c0.paragraphs[0].add_run(label_a); _set_run_font(r0, '黑体', 12, True)
        r1 = c1.paragraphs[0].add_run(label_b); _set_run_font(r1, '黑体', 12, True)
    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(7.5)
        row.cells[1].width = Cm(7.5)


# ============== 附件区 ==============

def _add_appendix(doc, items: List[str] = None):
    if not items:
        items = ["附件一: 合同正本", "附件二: 补充协议(如适用)", "附件三: 双方资质证明"]
    doc.add_page_break()
    _add_para(doc, "附 件 目 录", level=0, align='center', first_line_indent=False)
    _add_blank(doc, 1)
    for i, item in enumerate(items, 1):
        _add_para(doc, f"{i}. {item}", level=4, first_line_indent=False)


# ============== 修改对照表 ==============

def _add_change_table(doc, suggestions: List[Dict]):
    """修改对照表 - 修改版专用"""
    from docx.shared import Cm
    doc.add_page_break()
    _add_para(doc, "合 同 修 改 对 照 表", level=0, align='center', first_line_indent=False)
    _add_blank(doc, 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    # 表头
    headers = ["序号", "条款", "原内容摘要", "修改建议", "法律依据"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ''
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, '黑体', 11, True)
        # 灰色背景
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    for i, sug in enumerate(suggestions, 1):
        row = table.add_row()
        # 0: 序号
        c0 = row.cells[0]; c0.paragraphs[0].text = ''
        r = c0.paragraphs[0].add_run(str(i)); _set_run_font(r, '宋体', 10)
        # 1: 条款
        c1 = row.cells[1]; c1.paragraphs[0].text = ''
        r = c1.paragraphs[0].add_run(sug.get("clause", "")[:20]); _set_run_font(r, '宋体', 10)
        # 2: 原内容摘要
        c2 = row.cells[2]; c2.paragraphs[0].text = ''
        r = c2.paragraphs[0].add_run(sug.get("original_text", sug.get("content", ""))[:60]); _set_run_font(r, '宋体', 10)
        # 3: 修改建议
        c3 = row.cells[3]; c3.paragraphs[0].text = ''
        r = c3.paragraphs[0].add_run(sug.get("suggested_text", sug.get("suggestion", ""))[:60]); _set_run_font(r, '宋体', 10)
        # 4: 法律依据
        c4 = row.cells[4]; c4.paragraphs[0].text = ''
        r = c4.paragraphs[0].add_run(sug.get("legal_basis", "")[:40]); _set_run_font(r, '宋体', 10)
    # 列宽
    widths_cm = [1.0, 2.5, 4.5, 4.5, 2.5]
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


# ============== 批注框 (Word Comments) ==============

def _add_comments_part(doc):
    """为 docx 文档添加 comments part (存储所有批注)"""
    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap, qn
    
    # 检查是否已存在
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return rel.target_part
    
    # 创建 comments part
    partname = PackURI('/word/comments.xml')
    content_type = CT.WML_COMMENTS
    
    # 空 comments 根
    comments = OxmlElement('w:comments')
    blob = __import__('lxml.etree', fromlist=['tostring']).tostring(
        comments, xml_declaration=True, encoding='UTF-8', standalone=True
    )
    comments_part = Part(partname, content_type, blob, doc.part.package)
    doc.part.relate_to(comments_part, RT.COMMENTS)
    return comments_part


def _add_comment(doc, comments_part, comment_id: int, author: str, initials: str, text: str):
    """在 comments.xml 中添加一条批注"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsmap
    from lxml import etree
    
    # 解析现有 comments.xml
    root = etree.fromstring(comments_part.blob)
    
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), str(comment_id))
    comment.set(qn('w:author'), author)
    comment.set(qn('w:initials'), initials)
    comment.set(qn('w:date'), '2026-06-17T10:00:00Z')
    
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'CommentText')
    pPr.append(pStyle)
    p.append(pPr)
    
    # 风险等级标签
    r0 = OxmlElement('w:r')
    rPr0 = OxmlElement('w:rPr')
    rFonts0 = OxmlElement('w:rFonts')
    rFonts0.set(qn('w:eastAsia'), '黑体')
    rFonts0.set(qn('w:ascii'), 'Times New Roman')
    rPr0.append(rFonts0)
    b0 = OxmlElement('w:b'); rPr0.append(b0)
    sz0 = OxmlElement('w:sz'); sz0.set(qn('w:val'), '20'); rPr0.append(sz0)
    color0 = OxmlElement('w:color'); color0.set(qn('w:val'), 'C00000'); rPr0.append(color0)
    r0.append(rPr0)
    t0 = OxmlElement('w:t')
    t0.set(qn('xml:space'), 'preserve')
    t0.text = f'【律师批注】\n'
    r0.append(t0)
    p.append(r0)
    
    # 批注内容
    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    rFonts1 = OxmlElement('w:rFonts')
    rFonts1.set(qn('w:eastAsia'), '宋体')
    rFonts1.set(qn('w:ascii'), 'Times New Roman')
    rPr1.append(rFonts1)
    sz1 = OxmlElement('w:sz'); sz1.set(qn('w:val'), '20'); rPr1.append(sz1)
    r1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.set(qn('xml:space'), 'preserve')
    t1.text = text
    r1.append(t1)
    p.append(r1)
    
    comment.append(p)
    root.append(comment)
    
    # 写回
    new_blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    comments_part._blob = new_blob


def _wrap_run_with_comment(paragraph, run, comment_id: int):
    """在 run 前后插入批注范围标记 + 引用"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    r_elem = run._element
    parent = r_elem.getparent()
    idx = list(parent).index(r_elem)
    
    # commentRangeStart 放在 run 前
    crs = OxmlElement('w:commentRangeStart')
    crs.set(qn('w:id'), str(comment_id))
    parent.insert(idx, crs)
    
    # commentRangeEnd + commentReference 放在 run 后
    cre = OxmlElement('w:commentRangeEnd')
    cre.set(qn('w:id'), str(comment_id))
    parent.insert(idx + 2, cre)
    
    crr_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'CommentReference')
    rPr.append(rStyle)
    crr_run.append(rPr)
    crr = OxmlElement('w:commentReference')
    crr.set(qn('w:id'), str(comment_id))
    crr_run.append(crr)
    parent.insert(idx + 3, crr_run)


def _add_modified_paragraph(doc, text, level, comment_id=None, comments_part=None, lawyer='经办律师', risk_label='', reason='', suggestion='', is_modified_clause=False):
    """添加段落 + 可选批注框 + 建议文字颜色标注"""
    p = _add_para(doc, text, level=level)
    # 修改标记: [已修改] 条款的标题用深蓝色 (整体标记)
    # 具体修改的文字由 **加粗** → 蓝色, ~~删除~~ → 红色 处理
    if comment_id is not None and comments_part is not None:
        # 拼批注内容
        comment_text = f"风险等级: {_risk_label(risk_label)}\n"
        comment_text += f"问题: {reason[:200]}\n"
        if suggestion:
            comment_text += f"建议: {suggestion[:200]}\n"
        comment_text += "— " + lawyer
        _add_comment(doc, comments_part, comment_id, lawyer, lawyer[:1], comment_text)
        # 给段落里第一个 run 加批注包裹
        for run in p.runs:
            if run.text and run.text.strip():
                _wrap_run_with_comment(p, run, comment_id)
                break
        # 在段落后添加绿色建议文字
        if suggestion:
            from docx.shared import RGBColor, Pt
            p2 = doc.add_paragraph()
            p2.paragraph_format.first_line_indent = None
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(6)
            run_label = p2.add_run("【修改建议】")
            _set_run_font(run_label, '宋体', 10, bold=True, color=RGBColor(0x00, 0x80, 0x00))
            run_sug = p2.add_run(suggestion[:300])
            _set_run_font(run_sug, '宋体', 10, color=RGBColor(0x00, 0x80, 0x00))


# ============== 修改版主函数 ==============

def generate_modified_docx(
    content: str,
    suggestions: List[Dict],
    contract_title: str = "合同",
    contract_no: str = "",
    lawyer: str = DEFAULT_LAWYER,
    date_str: str = DEFAULT_DATE,
    risk_level: str = ""
) -> bytes:
    """生成修改版 docx
    content: 原合同正文 (纯文本或 markdown)
    suggestions: 修改建议列表 [{clause, original_text, suggested_text, reason, legal_basis, risk_level}]
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_BREAK
    
    if not contract_no:
        contract_no = f"{CONTRACT_NO_PREFIX}-{abs(hash(contract_title)) % 10000:04d}"
    
    doc = Document()
    _setup_docx_default_style(doc)
    _setup_header_footer(doc, contract_title, contract_no)
    
    # 初始化批注 part
    comments_part = _add_comments_part(doc)
    
    # 正文 - 智能插入批注
    # 匹配意见到原文位置
    matched_indices = set()
    lines = content.split('\n')
    comment_id = 0
    from docx.shared import RGBColor
    BLUE = RGBColor(0x00, 0x66, 0xCC)
    
    in_mod = False  # 是否在修改条款范围内
    mod_lvl = -1
    cur_sug_text = None  # 当前条款的修改建议
    
    for line in lines:
        s = line.strip()
        if not s:
            # 空行: 如果在修改范围且有建议, 先输出修改说明
            if in_mod and cur_sug_text:
                _add_para(doc, f"【修改说明】{cur_sug_text}", level=9, color=BLUE)
                cur_sug_text = None
            _add_blank(doc, 1)
            continue
        s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
        s_clean = re.sub(r'`([^`]+)`', r'\1', s)
        is_modified_clause = '[已修改]' in s_clean
        s_clean = s_clean.replace(' [已修改]', '').replace('[已修改]', '').strip()
        level, cleaned = _detect_clause_level(s_clean)
        
        # 退出修改范围: 遇到同级或更高级标题 (level <= mod_lvl 且不是子项)
        if in_mod and level >= 0 and level <= mod_lvl:
            if cur_sug_text:
                _add_para(doc, f"【修改说明】{cur_sug_text}", level=9, color=BLUE)
                cur_sug_text = None
            in_mod = False
            mod_lvl = -1
        
        if is_modified_clause:
            in_mod = True
            mod_lvl = level
        
        # 按顺序匹配 suggestion 到 [已修改] 条款
        match = None
        if is_modified_clause and matched_indices:
            pass  # 已有匹配的 suggestion 在处理中
        if is_modified_clause:
            # 找下一个未匹配的 suggestion
            for idx, sug in enumerate(suggestions):
                if idx not in matched_indices:
                    match = (idx, sug)
                    break
        
        should_blue = is_modified_clause or in_mod
        
        if match:
            idx, sug = match
            matched_indices.add(idx)
            comment_id += 1
            cur_sug_text = sug.get("suggested_text") or sug.get("suggestion", "")
            _add_modified_paragraph(
                doc, cleaned, level=level,
                comment_id=comment_id, comments_part=comments_part,
                lawyer=lawyer,
                risk_label=sug.get("risk_level", "medium"),
                reason=sug.get("reason") or sug.get("content", ""),
                suggestion=sug.get("suggested_text") or sug.get("suggestion", ""),
                is_modified_clause=should_blue
            )
        elif should_blue:
            p = _add_para(doc, cleaned, level=level, color=BLUE)
        else:
            _add_para(doc, cleaned, level=level)
    
    # 文件末尾还在修改范围内
    if in_mod and cur_sug_text:
        _add_para(doc, f"【修改说明】{cur_sug_text}", level=9, color=BLUE)
    
    # 4) 未匹配的意见 - 智能匹配或集中说明
    # 如果 AI 已经将修改融入正文 (没有 [已修改] 标记), 则用关键词匹配
    # 如果仍无法匹配且正文已包含建议内容, 则跳过 (说明已融入)
    import unicodedata as _uc
    def _normalize_text(t):
        """标准化文本: 去掉引号、空格等差异"""
        if not t:
            return ''
        t = t.replace('\u2018', '').replace('\u2019', '')  # 去掉弯引号
        t = t.replace('\u201c', '').replace('\u201d', '')  # 去掉弯双引号
        t = t.replace("'", '').replace('"', '')  # 去掉直引号
        t = t.replace('\u3000', ' ').replace('\xa0', ' ')  # 全角空格
        t = re.sub(r'[\s\u3000\xa0]+', '', t)  # 去所有空格
        t = t.strip('\u3002\uff0c\uff1b\uff1a')  # 去结尾标点
        return t.strip()
    
    norm_content = _normalize_text(content)
    has_modification_markers = '[已修改]' in content
    unmatched = []
    for i, sug in enumerate(suggestions):
        if i in matched_indices:
            continue
        # 如果没有 [已修改] 标记, 说明 AI 直接融入了修改, 检查建议是否已在正文中
        if not has_modification_markers:
            sug_text = (sug.get('suggested_text') or sug.get('suggestion') or '').strip()
            if sug_text and len(sug_text) > 5:
                # 标准化后取前30个字符匹配
                norm_sug = _normalize_text(sug_text[:30])
                if norm_sug and norm_sug in norm_content:
                    matched_indices.add(i)
                    continue
        unmatched.append(sug)
    
    if not has_modification_markers:
        # AI 直接融入模式: 再检查剩余 unmatched
        truly_unmatched = []
        for sug in unmatched:
            sug_text = (sug.get('suggested_text') or sug.get('suggestion') or '').strip()
            if sug_text and len(sug_text) > 5:
                norm_sug = _normalize_text(sug_text[:30])
                if norm_sug and norm_sug in norm_content:
                    continue  # 已融入正文
            truly_unmatched.append(sug)
        unmatched = truly_unmatched
    
    if unmatched:
        doc.add_page_break()
        _add_para(doc, "其 他 修 改 建 议", level=0, align='center', first_line_indent=False)
        _add_blank(doc, 1)
        _add_para(doc, "以下修改建议不直接对应合同原文某一条款, 请在合同全文中全面检查：", level=9)
        _add_blank(doc, 1)
        for i, sug in enumerate(unmatched, 1):
            comment_id += 1
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.line_spacing_rule = 3
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(f"{i}. 【{sug.get('clause', '一般条款')}】 ")
            _set_run_font(run, '黑体', 12, True)
            run2 = p.add_run(sug.get('content', sug.get('reason', ''))[:200])
            _set_run_font(run2, '宋体', 12)
            # 加批注
            ctext = f"【{_risk_label(sug.get('risk_level', 'medium'))}】\n"
            ctext += f"问题: {sug.get('content', sug.get('reason', ''))[:200]}\n"
            if sug.get('suggested_text') or sug.get('suggestion'):
                ctext += f"建议: {(sug.get('suggested_text') or sug.get('suggestion', ''))[:200]}\n"
            ctext += "— " + lawyer
            _add_comment(doc, comments_part, comment_id, lawyer, lawyer[:1], ctext)
            _wrap_run_with_comment(p, run, comment_id)
            if sug.get('suggested_text') or sug.get('suggestion'):
                _add_para(doc, f"建议: {(sug.get('suggested_text') or sug.get('suggestion', ''))[:200]}", level=9, color=RGBColor(0x00, 0x66, 0x00))
    
    # 5) 修改对照表
    _add_change_table(doc, suggestions)
    
    # 6) 签字盖章
    _add_signature_block(doc)
    
    # 7) 附件
    _add_appendix(doc)
    
    # 保存
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============== 清洁版主函数 ==============

def _clean_export_content(content: str) -> str:
    """清洗导出内容: 去掉标记、修改说明、markdown代码块等"""
    c = content
    c = re.sub(r'^```(?:markdown|md)?\s*\n?', '', c)
    c = re.sub(r'\n?```\s*$', '', c)
    c = re.sub(r'\s*\[已修改\]', '', c)
    c = re.sub(r'【修改说明】[^\n]*\n?', '', c)
    c = re.sub(r'---+\s*修改说明.*?(?=\n\n|\Z)', '', c, flags=re.DOTALL)
    c = re.sub(r'^修改说明\s*\n(?:\d+\..*?\n)+', '', c, flags=re.MULTILINE)
    c = re.sub(r'以下修改建议不直接对应.*?(?=\n\n|\Z)', '', c, flags=re.DOTALL)
    return c


def generate_clean_docx(
    content: str,
    contract_title: str = "合同",
    contract_no: str = "",
    lawyer: str = DEFAULT_LAWYER,
    date_str: str = DEFAULT_DATE,
    risk_level: str = ""
) -> bytes:
    """生成清洁版 docx - 应用 AI 建议后重排, 无痕迹"""
    from docx import Document
    from docx.shared import Cm
    
    if not contract_no:
        contract_no = f"{CONTRACT_NO_PREFIX}-{abs(hash(contract_title)) % 10000:04d}"
    
    # 清洗内容
    clean_content = _clean_export_content(content)
    
    doc = Document()
    _setup_docx_default_style(doc)
    _setup_header_footer(doc, contract_title, contract_no)
    
    # 正文 - 无痕迹
    _render_body(doc, clean_content)
    
    # 签字盖章
    _add_signature_block(doc)
    
    # 附件
    _add_appendix(doc)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============== 原始版主函数 ==============

def generate_original_docx(
    content: str,
    contract_title: str = "合同",
    contract_no: str = "",
    lawyer: str = DEFAULT_LAWYER,
    date_str: str = DEFAULT_DATE
) -> bytes:
    """生成原始版 docx - 原文件重排律所级版式, 但不改任何内容"""
    return generate_clean_docx(content, contract_title, contract_no, lawyer, date_str, risk_level="")


# ============== PDF 导出 (基于 reportlab) ==============

def _register_cjk():
    """注册中文字体到 reportlab"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    candidates = [
        # macOS
        ('CJK', '/System/Library/Fonts/STHeiti Light.ttc', 0),
        ('CJK', '/System/Library/Fonts/PingFang.ttc', 0),
        ('CJK', '/Library/Fonts/Songti.ttc', 0),
        ('CJK', '/System/Library/Fonts/Supplemental/Songti.ttc', 0),
        ('CJK', '/System/Library/Fonts/Supplemental/STSong.ttf', 0),
        # Linux
        ('CJK', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc', 0),
        ('CJK', '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc', 0),
        ('CJK', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc', 0),
        ('CJK', '/usr/share/fonts/truetype/arphic/uming.ttc', 0),
    ]
    for name, path, sub in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=sub))
                return name
            except Exception:
                continue
    # 退化: Helvetica
    return 'Helvetica'


def _build_pdf_styles(cjk_font: str):
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    base = getSampleStyleSheet()['Normal']
    return {
        'title': ParagraphStyle('Title', parent=base, fontName=cjk_font, fontSize=22, leading=33, alignment=TA_CENTER, spaceAfter=24),
        'h1': ParagraphStyle('H1', parent=base, fontName=cjk_font, fontSize=16, leading=24, spaceBefore=18, spaceAfter=8, textColor=colors.black),
        'h2': ParagraphStyle('H2', parent=base, fontName=cjk_font, fontSize=14, leading=21, spaceBefore=10, spaceAfter=6),
        'h3': ParagraphStyle('H3', parent=base, fontName=cjk_font, fontSize=12, leading=18, spaceBefore=6, spaceAfter=4),
        'body': ParagraphStyle('Body', parent=base, fontName=cjk_font, fontSize=11, leading=18, firstLineIndent=22, alignment=TA_JUSTIFY),
        'body_no_indent': ParagraphStyle('BodyNoIndent', parent=base, fontName=cjk_font, fontSize=11, leading=18),
        'footer': ParagraphStyle('Footer', parent=base, fontName=cjk_font, fontSize=8, leading=10, textColor=colors.grey, alignment=TA_CENTER),
    }


def _draw_page_decoration(canvas, doc, contract_no: str, contract_title: str, cjk_font: str = 'CJK'):
    """页眉页脚 - 通用商务版式 (中文字体)"""
    from reportlab.lib.units import cm
    canvas.saveState()
    # 页眉
    canvas.setFont(cjk_font, 9)
    canvas.setFillColorRGB(0.3, 0.3, 0.3)
    canvas.drawString(2 * cm, 27.5 * cm, contract_title)
    canvas.drawRightString(19 * cm, 27.5 * cm, f"第 {doc.page} 页")
    # 品牌色页眉线
    canvas.setStrokeColorRGB(0.12, 0.22, 0.39)
    canvas.setLineWidth(1.0)
    canvas.line(2 * cm, 27.3 * cm, 19 * cm, 27.3 * cm)
    # 页脚
    canvas.setFont(cjk_font, 8)
    canvas.setFillColorRGB(0.5, 0.5, 0.5)
    canvas.drawString(2 * cm, 1 * cm, f"编号: {contract_no}")
    canvas.drawRightString(19 * cm, 1 * cm, "合同审查报告")
    canvas.restoreState()


def generate_pdf_from_docx_args(
    content: str,
    contract_title: str = "合同",
    contract_no: str = "",
    suggestions: Optional[List[Dict]] = None,
    is_modified: bool = False
) -> bytes:
    """PDF 通用生成: clean / original / modified 都能用"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    
    if not contract_no:
        contract_no = f"{CONTRACT_NO_PREFIX}-{abs(hash(contract_title)) % 10000:04d}"
    
    cjk = _register_cjk()
    styles = _build_pdf_styles(cjk)
    
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3.18 * cm, rightMargin=2.54 * cm,
        topMargin=2.54 * cm, bottomMargin=2.54 * cm,
        title=contract_title,
        author="合同审查系统",
    )
    
    story = []
    in_mod = False
    mod_lvl = -1
    cur_sug = None
    
    pdf_lines = content.split('\n')
    i = 0
    while i < len(pdf_lines):
        s = pdf_lines[i].strip()
        if not s:
            if in_mod and cur_sug:
                story.append(Paragraph(f'<font color="#0066CC">\u3010\u4fee\u6539\u8bf4\u660e\u3011{cur_sug}</font>', styles['body']))
                cur_sug = None
            story.append(Spacer(1, 0.5 * cm))
            i += 1
            continue
        
        # 检测 Markdown 表格
        if _is_md_table_row(s):
            table_lines = []
            while i < len(pdf_lines) and _is_md_table_row(pdf_lines[i].strip()):
                row_s = pdf_lines[i].strip()
                if not _is_md_table_separator(row_s):
                    table_lines.append(_parse_md_table_row(row_s))
                i += 1
            if len(table_lines) >= 2:
                # 构建 reportlab Table
                from reportlab.lib import colors as rl_colors
                # 清理单元格文本
                clean_rows = []
                for ri, row in enumerate(table_lines):
                    clean_cells = []
                    for cell in row:
                        c = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
                        c = re.sub(r'`([^`]+)`', r'\1', c)
                        c = c.replace('<', '&lt;').replace('>', '&gt;')
                        clean_cells.append(c)
                    clean_rows.append(clean_cells)
                num_cols = len(clean_rows[0])
                # 补齐列数
                for r in clean_rows:
                    while len(r) < num_cols:
                        r.append('')
                # 转为 Paragraph 以支持换行
                para_rows = []
                for ri, row in enumerate(clean_rows):
                    para_cells = []
                    for ci, cell_text in enumerate(row[:num_cols]):
                        if ri == 0:
                            p = Paragraph(f'<b>{cell_text}</b>', styles['body_no_indent'])
                        else:
                            p = Paragraph(cell_text, styles['body_no_indent'])
                        para_cells.append(p)
                    para_rows.append(para_cells)
                # 计算列宽
                avail_width = A4[0] - 3.18 * cm - 2.54 * cm
                col_w = avail_width / num_cols
                pdf_table = Table(para_rows, colWidths=[col_w] * num_cols)
                pdf_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#D9E2F3')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 0.5 * cm))
            continue
        
        s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
        s = re.sub(r'`([^`]+)`', r'\1', s)
        is_mod = '[已修改]' in s
        s = s.replace(' [已修改]', '').replace('[已修改]', '').strip()
        level, cleaned = _detect_clause_level(s)
        
        if in_mod and level >= 0 and level <= mod_lvl:
            if cur_sug:
                story.append(Paragraph(f'<font color="#0066CC">\u3010\u4fee\u6539\u8bf4\u660e\u3011{cur_sug}</font>', styles['body']))
                cur_sug = None
            in_mod = False
            mod_lvl = -1
        
        if is_mod:
            in_mod = True
            mod_lvl = level
        
        show_blue = is_mod or in_mod
        mc = '#0066CC' if show_blue else None
        
        def _para(text, style, _mc=mc):
            if _mc:
                return Paragraph(f'<font color="{_mc}">{text}</font>', style)
            return Paragraph(text, style)
        
        try:
            if level == 0:
                story.append(_para(cleaned, styles['title']))
            elif level == 1:
                story.append(_para(cleaned, styles['h1']))
            elif level == 2:
                story.append(_para(cleaned, styles['h2']))
            elif level == 3:
                story.append(_para(cleaned, styles['h3']))
            elif level == 4 or level == 5:
                story.append(_para(cleaned, styles['body_no_indent']))
            else:
                story.append(_para(cleaned, styles['body']))
        except Exception:
            cleaned_safe = cleaned.replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(cleaned_safe, styles['body']))
        i += 1
    
    if in_mod and cur_sug:
        story.append(Paragraph(f'<font color="#0066CC">\u3010\u4fee\u6539\u8bf4\u660e\u3011{cur_sug}</font>', styles['body']))

    # 修改版: 末尾加修改建议
    if is_modified and suggestions:
        story.append(PageBreak())
        story.append(Paragraph("修改建议汇总", styles['title']))
        story.append(Spacer(1, 0.5 * cm))
        for i, sug in enumerate(suggestions, 1):
            clause = sug.get('clause', '一般条款')
            reason = sug.get('reason') or sug.get('content', '')
            suggestion = sug.get('suggested_text') or sug.get('suggestion', '')
            risk = sug.get('risk_level', 'medium')
            risk_color = {'high': '#FF0000', 'medium': '#FF8C00', 'low': '#228B22'}.get(risk, '#FF8C00')
            story.append(Paragraph(f"{i}. 【{clause}】", styles['h2']))
            story.append(Paragraph(f'<font color="{risk_color}">风险等级: {risk}</font>', styles['body']))
            story.append(Paragraph(f"问题: {reason[:200]}", styles['body']))
            if suggestion:
                story.append(Paragraph(f'<font color="#006600"><b>建议:</b> {suggestion[:200]}</font>', styles['body']))
            story.append(Spacer(1, 0.3 * cm))

    pdf.build(story, onFirstPage=lambda c, d: _draw_page_decoration(c, d, contract_no, contract_title, cjk),
              onLaterPages=lambda c, d: _draw_page_decoration(c, d, contract_no, contract_title, cjk))
    buf.seek(0)
    return buf.getvalue()


def generate_clean_pdf(content: str, contract_title: str = "合同", contract_no: str = "", **kw) -> bytes:
    """生成清洁版 PDF - 无痕迹"""
    return generate_pdf_from_docx_args(_clean_export_content(content), contract_title, contract_no, is_modified=False)


def generate_original_pdf(content: str, contract_title: str = "合同", contract_no: str = "", **kw) -> bytes:
    """生成原始版 PDF"""
    return generate_pdf_from_docx_args(_clean_export_content(content), contract_title, contract_no, is_modified=False)


def generate_modified_pdf(content: str, suggestions: List[Dict], contract_title: str = "合同", contract_no: str = "", **kw) -> bytes:
    """生成修改版 PDF - 含修改建议"""
    return generate_pdf_from_docx_args(content, contract_title, contract_no, suggestions=suggestions, is_modified=True)
