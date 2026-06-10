"""
审查报告导出服务
支持Word和PDF格式导出
"""

import io
import json
import logging
from datetime import datetime
from typing import Dict, Optional

logger = logging.getLogger(__name__)


def generate_word_report(review_data: dict, contract_data: dict) -> bytes:
    """生成Word格式审查报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.warning("python-docx not installed, using fallback")
        return _generate_text_report(review_data, contract_data).encode("utf-8")

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)

    # 标题
    title = doc.add_heading("合同审查报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表
    doc.add_heading("一、合同基本信息", level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"

    info_rows = [
        ("合同名称", contract_data.get("title", "-")),
        ("合同类型", contract_data.get("contract_type", "-")),
        ("甲方", contract_data.get("party_a", "-")),
        ("乙方", contract_data.get("party_b", "-")),
        ("合同金额", f"{contract_data.get('amount', '-')} {contract_data.get('currency', 'CNY')}"),
        ("签署日期", str(contract_data.get("sign_date", "-"))),
        ("有效期", f"{contract_data.get('effective_date', '-')} 至 {contract_data.get('expiry_date', '-')}"),
        ("审查时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    ]

    for i, (label, value) in enumerate(info_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # 加粗标签列
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 风险评估总览
    doc.add_heading("二、风险评估总览", level=1)

    risk_level = review_data.get("risk_level", "unknown")
    risk_score = review_data.get("risk_score", 0)
    risk_color = {"high": "红色", "medium": "橙色", "low": "绿色"}.get(risk_level, "灰色")

    p = doc.add_paragraph()
    p.add_run(f"综合风险等级：").bold = True
    run = p.add_run(f"【{risk_level.upper()}】")
    run.font.color.rgb = RGBColor(255, 0, 0) if risk_level == "high" else (
        RGBColor(255, 165, 0) if risk_level == "medium" else RGBColor(0, 128, 0)
    )
    run.bold = True

    p = doc.add_paragraph()
    p.add_run(f"风险评分：").bold = True
    p.add_run(f"{risk_score}/100")

    p = doc.add_paragraph()
    p.add_run(f"审查摘要：").bold = True
    p.add_run(review_data.get("summary", "无"))

    # 多Agent审查结果
    agent_results = review_data.get("agent_results", {})
    if agent_results:
        doc.add_heading("三、多维度审查详情", level=1)

        for agent_id, result in agent_results.items():
            doc.add_heading(f"{result.get('icon', '')} {result.get('agent_name', agent_id)}", level=2)

            p = doc.add_paragraph()
            p.add_run(f"审查重点：").bold = True
            p.add_run(result.get("focus", "-"))

            p = doc.add_paragraph()
            p.add_run(f"风险评分：").bold = True
            p.add_run(f"{result.get('risk_score', '-')}/100")

            p = doc.add_paragraph()
            p.add_run(f"摘要：").bold = True
            p.add_run(result.get("summary", "无"))

            findings = result.get("findings", [])
            if findings:
                doc.add_paragraph(f"发现 {len(findings)} 项问题：")
                for j, finding in enumerate(findings, 1):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{j}. {finding.get('title', '')}").bold = True
                    if finding.get("description"):
                        doc.add_paragraph(f"   描述：{finding['description']}")
                    if finding.get("suggestion"):
                        doc.add_paragraph(f"   建议：{finding['suggestion']}")

    # 风控规则检查
    industry_risks = review_data.get("industry_risks", [])
    if industry_risks:
        doc.add_heading("四、行业风控规则检查", level=1)
        table = doc.add_table(rows=len(industry_risks) + 1, cols=4)
        table.style = "Table Grid"
        headers = ["规则编号", "规则名称", "严重度", "建议"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True

        for i, risk in enumerate(industry_risks, 1):
            table.rows[i].cells[0].text = risk.get("rule_id", "")
            table.rows[i].cells[1].text = risk.get("rule_name", "")
            table.rows[i].cells[2].text = f"{risk.get('severity', 0):.0%}"
            table.rows[i].cells[3].text = risk.get("suggestion", "")

    # 毒丸条款
    poison_pills = review_data.get("poison_pills", [])
    if poison_pills:
        doc.add_heading("五、毒丸条款检测", level=1)
        for pp in poison_pills:
            p = doc.add_paragraph()
            p.add_run(f"⚠️ {pp.get('name', '')}").bold = True
            doc.add_paragraph(f"   类型：{pp.get('type', '')} | 严重度：{pp.get('severity', 0):.0%}")
            doc.add_paragraph(f"   匹配内容：{pp.get('matched_text', '')[:200]}")

    # 签批建议
    doc.add_heading("六、签批建议", level=1)
    if risk_level == "high":
        doc.add_paragraph("❌ 建议：退回修改，待高风险问题解决后重新提交审查")
    elif risk_level == "medium":
        doc.add_paragraph("⚠️ 建议：附条件批准，要求在签署前解决以下问题")
    else:
        doc.add_paragraph("✅ 建议：同意签署，风险在可控范围内")

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_pdf_report(review_data: dict, contract_data: dict) -> bytes:
    """生成PDF格式审查报告"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        logger.warning("reportlab not installed, generating text report")
        return _generate_text_report(review_data, contract_data).encode("utf-8")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)

    styles = getSampleStyleSheet()
    # 中文样式
    try:
        pdfmetrics.registerFont(TTFont("SimSun", "/usr/share/fonts/truetype/simsun.ttc"))
        normal_style = ParagraphStyle("Chinese", parent=styles["Normal"], fontName="SimSun", fontSize=10)
        heading_style = ParagraphStyle("ChineseHeading", parent=styles["Heading1"], fontName="SimSun", fontSize=14)
    except Exception:
        normal_style = styles["Normal"]
        heading_style = styles["Heading1"]

    elements = []

    # 标题
    elements.append(Paragraph("合同审查报告", heading_style))
    elements.append(Spacer(1, 0.5*cm))

    # 基本信息
    elements.append(Paragraph("一、合同基本信息", heading_style))
    info_data = [
        ["合同名称", contract_data.get("title", "-")],
        ["合同类型", contract_data.get("contract_type", "-")],
        ["甲方", contract_data.get("party_a", "-")],
        ["乙方", contract_data.get("party_b", "-")],
        ["合同金额", f"{contract_data.get('amount', '-')} {contract_data.get('currency', 'CNY')}"],
        ["审查时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
    ]
    t = Table(info_data, colWidths=[4*cm, 12*cm])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (0,-1), colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.5*cm))

    # 风险评估
    elements.append(Paragraph("二、风险评估总览", heading_style))
    risk_level = review_data.get("risk_level", "unknown")
    risk_score = review_data.get("risk_score", 0)
    elements.append(Paragraph(f"综合风险等级：【{risk_level.upper()}】 评分：{risk_score}/100", normal_style))
    elements.append(Paragraph(f"审查摘要：{review_data.get('summary', '无')}", normal_style))
    elements.append(Spacer(1, 0.5*cm))

    # Agent审查结果
    agent_results = review_data.get("agent_results", {})
    if agent_results:
        elements.append(Paragraph("三、多维度审查详情", heading_style))
        for agent_id, result in agent_results.items():
            elements.append(Paragraph(f"{result.get('icon', '')} {result.get('agent_name', agent_id)}", heading_style))
            elements.append(Paragraph(f"风险评分：{result.get('risk_score', '-')}/100", normal_style))
            elements.append(Paragraph(f"摘要：{result.get('summary', '无')}", normal_style))
            for finding in result.get("findings", []):
                elements.append(Paragraph(f"• {finding.get('title', '')}: {finding.get('description', '')[:100]}", normal_style))
            elements.append(Spacer(1, 0.3*cm))

    # 签批建议
    elements.append(Paragraph("四、签批建议", heading_style))
    if risk_level == "high":
        elements.append(Paragraph("❌ 建议：退回修改", normal_style))
    elif risk_level == "medium":
        elements.append(Paragraph("⚠️ 建议：附条件批准", normal_style))
    else:
        elements.append(Paragraph("✅ 建议：同意签署", normal_style))

    doc.build(elements)
    return buffer.getvalue()


def _generate_text_report(review_data: dict, contract_data: dict) -> str:
    """纯文本报告（降级方案）"""
    lines = [
        "=" * 60,
        "合同审查报告",
        "=" * 60,
        "",
        f"合同名称：{contract_data.get('title', '-')}",
        f"合同类型：{contract_data.get('contract_type', '-')}",
        f"甲方：{contract_data.get('party_a', '-')}",
        f"乙方：{contract_data.get('party_b', '-')}",
        f"合同金额：{contract_data.get('amount', '-')} {contract_data.get('currency', 'CNY')}",
        f"审查时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "-" * 60,
        f"风险等级：{review_data.get('risk_level', 'unknown').upper()}",
        f"风险评分：{review_data.get('risk_score', 0)}/100",
        f"摘要：{review_data.get('summary', '无')}",
        "-" * 60,
        "",
    ]

    for agent_id, result in review_data.get("agent_results", {}).items():
        lines.append(f"{result.get('icon', '')} {result.get('agent_name', agent_id)}")
        lines.append(f"  评分：{result.get('risk_score', '-')}/100")
        lines.append(f"  摘要：{result.get('summary', '无')}")
        for finding in result.get("findings", []):
            lines.append(f"  • {finding.get('title', '')}: {finding.get('description', '')[:80]}")
        lines.append("")

    lines.append("=" * 60)
    risk_level = review_data.get("risk_level", "unknown")
    if risk_level == "high":
        lines.append("签批建议：❌ 退回修改")
    elif risk_level == "medium":
        lines.append("签批建议：⚠️ 附条件批准")
    else:
        lines.append("签批建议：✅ 同意签署")

    return "\n".join(lines)
