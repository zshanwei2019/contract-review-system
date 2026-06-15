from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, BackgroundTasks, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from sqlalchemy.orm import selectinload
from typing import Optional, List
from dataclasses import asdict
import os
import uuid
import json
from datetime import datetime

from app.core.database import get_db
from app.core.security import get_current_user, require_role
from app.core.config import settings
from app.models.user import User
from app.models.contract import Contract, ContractType, ContractStatus, ContractVersion, ContractFile
from app.models.review import ReviewTask, ReviewTaskStatus, ReviewOpinion
from app.schemas.contract import (
    ContractCreate,
    ContractUpdate,
    ContractResponse,
    ContractList,
    ContractVersionResponse,
    ContractFileResponse,
    ContractUploadResponse,
)
from app.services.ai_review import review_contract_with_ai, extract_file_content

router = APIRouter()


def generate_contract_no() -> str:
    """Generate contract number"""
    now = datetime.now()
    return f"HT-{now.strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"


@router.get("", response_model=ContractList)
async def list_contracts(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    keyword: Optional[str] = None,
    contract_type: Optional[ContractType] = None,
    status: Optional[ContractStatus] = None,
    department: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同列表"""
    query = select(Contract)
    count_query = select(func.count()).select_from(Contract)
    
    if keyword:
        query = query.where(
            (Contract.title.contains(keyword)) |
            (Contract.contract_no.contains(keyword)) |
            (Contract.party_a.contains(keyword)) |
            (Contract.party_b.contains(keyword))
        )
        count_query = count_query.where(
            (Contract.title.contains(keyword)) |
            (Contract.contract_no.contains(keyword)) |
            (Contract.party_a.contains(keyword)) |
            (Contract.party_b.contains(keyword))
        )
    
    if contract_type:
        query = query.where(Contract.contract_type == contract_type)
        count_query = count_query.where(Contract.contract_type == contract_type)
    
    if status:
        query = query.where(Contract.status == status)
        count_query = count_query.where(Contract.status == status)
    
    if department:
        query = query.where(Contract.department == department)
        count_query = count_query.where(Contract.department == department)
    
    # Get total count
    total_result = await db.execute(count_query)
    total = total_result.scalar()
    
    # Get paginated results
    query = query.order_by(Contract.created_at.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)
    result = await db.execute(query)
    contracts = result.scalars().all()
    
    return ContractList(total=total, items=contracts)


@router.post("", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
@router.post("", response_model=ContractResponse)
async def create_contract(
    title: str = Form(...),
    contract_type: ContractType = Form(ContractType.OTHER),
    party_a: Optional[str] = Form(None),
    party_b: Optional[str] = Form(None),
    amount: Optional[float] = Form(None),
    currency: Optional[str] = Form("CNY"),
    sign_date: Optional[str] = Form(None),
    effective_date: Optional[str] = Form(None),
    expiry_date: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    department: Optional[str] = Form(None),
    project_name: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """创建合同（支持同时上传文件）"""
    # 处理文件上传
    file_path = None
    file_name = None
    file_size = None
    file_type = None
    extracted_info = None
    
    if file:
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"不支持的文件格式，支持: {', '.join(settings.ALLOWED_EXTENSIONS)}",
            )
        
        content = await file.read()
        if len(content) > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"文件大小超过限制（最大{settings.MAX_FILE_SIZE // 1024 // 1024}MB）",
            )
        
        upload_dir = os.path.join(settings.UPLOAD_DIR, "contracts")
        os.makedirs(upload_dir, exist_ok=True)
        
        saved_name = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(upload_dir, saved_name)
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        file_name = file.filename
        file_size = len(content)
        file_type = file_ext
        
        # AI自动提取合同基本信息
        from app.services.ai_review import extract_contract_info
        extracted_info = await extract_contract_info(file_path)
    
    contract = Contract(
        contract_no=generate_contract_no(),
        title=title or (extracted_info.get("title") if extracted_info else title),
        contract_type=contract_type or (ContractType(extracted_info.get("contract_type")) if extracted_info and extracted_info.get("contract_type") else ContractType.OTHER),
        party_a=party_a or (extracted_info.get("party_a") if extracted_info else None),
        party_b=party_b or (extracted_info.get("party_b") if extracted_info else None),
        amount=amount or (extracted_info.get("amount") if extracted_info else None),
        currency=currency or (extracted_info.get("currency") if extracted_info else "CNY"),
        sign_date=datetime.strptime(sign_date, "%Y-%m-%d") if sign_date else (datetime.strptime(extracted_info["sign_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("sign_date") else None),
        effective_date=datetime.strptime(effective_date, "%Y-%m-%d") if effective_date else (datetime.strptime(extracted_info["effective_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("effective_date") else None),
        expiry_date=datetime.strptime(expiry_date, "%Y-%m-%d") if expiry_date else (datetime.strptime(extracted_info["expiry_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("expiry_date") else None),
        description=description or (extracted_info.get("description") if extracted_info else None),
        department=department,
        project_name=project_name,
        file_path=file_path,
        file_name=file_name,
        file_size=file_size,
        file_type=file_type,
        uploader_id=current_user.id,
        status=ContractStatus.DRAFT,
    )
    db.add(contract)
    await db.commit()
    await db.refresh(contract)
    
    # 如果上传了文件，自动保存文件记录
    if file_path and file_name:
        contract_file = ContractFile(
            contract_id=contract.id,
            file_name=file_name,
            file_path=file_path,
            file_size=file_size,
            file_type=file_type,
            uploader_id=current_user.id,
        )
        db.add(contract_file)
        await db.commit()
    
    return contract


@router.post("/upload", response_model=ContractUploadResponse)
async def upload_contract(
    file: UploadFile = File(...),
    title: Optional[str] = None,
    contract_type: Optional[ContractType] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """上传合同文件，自动识别基本信息"""
    # Validate file extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"不支持的文件格式，支持: {', '.join(settings.ALLOWED_EXTENSIONS)}",
        )
    
    # Validate file size
    file_content = await file.read()
    if len(file_content) > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"文件大小超过限制（最大{settings.MAX_FILE_SIZE // 1024 // 1024}MB）",
        )
    
    # Save file
    upload_dir = os.path.join(settings.UPLOAD_DIR, "contracts")
    os.makedirs(upload_dir, exist_ok=True)
    
    saved_name = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(upload_dir, saved_name)
    
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    # AI自动提取合同基本信息
    from app.services.file_parser import extract_text_from_file
    from app.services.ai_review import extract_contract_info
    
    extracted_info = await extract_contract_info(file_path)
    
    # 使用提取的信息或传入的参数
    contract_title = title or (extracted_info.get("title") if extracted_info else None) or file.filename
    contract_type_val = contract_type or (extracted_info.get("contract_type") if extracted_info else "other")
    
    # Create contract
    contract = Contract(
        contract_no=generate_contract_no(),
        title=contract_title,
        contract_type=ContractType(contract_type_val) if contract_type_val else ContractType.OTHER,
        party_a=extracted_info.get("party_a") if extracted_info else None,
        party_b=extracted_info.get("party_b") if extracted_info else None,
        amount=extracted_info.get("amount") if extracted_info else None,
        currency=extracted_info.get("currency") if extracted_info else "CNY",
        sign_date=datetime.strptime(extracted_info["sign_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("sign_date") else None,
        effective_date=datetime.strptime(extracted_info["effective_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("effective_date") else None,
        expiry_date=datetime.strptime(extracted_info["expiry_date"], "%Y-%m-%d") if extracted_info and extracted_info.get("expiry_date") else None,
        description=extracted_info.get("description") if extracted_info else None,
        file_path=file_path,
        file_name=file.filename,
        file_size=len(file_content),
        file_type=file_ext,
        uploader_id=current_user.id,
        status=ContractStatus.DRAFT,
    )
    db.add(contract)
    await db.commit()
    await db.refresh(contract)
    
    return ContractUploadResponse(
        id=contract.id,
        file_name=file.filename,
        file_path=file_path,
        status="success",
        message="合同上传成功，已自动识别基本信息",
    )


@router.get("/{contract_id}", response_model=ContractResponse)
async def get_contract(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同详情"""
    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    return contract


@router.put("/{contract_id}", response_model=ContractResponse)
async def update_contract(
    contract_id: int,
    contract_data: ContractUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """更新合同"""
    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    # Update fields
    update_data = contract_data.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(contract, field, value)
    
    await db.commit()
    await db.refresh(contract)
    
    return contract


@router.delete("/{contract_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_contract(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """删除合同"""
    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    # Delete file if exists
    if contract.file_path and os.path.exists(contract.file_path):
        os.remove(contract.file_path)
    
    await db.delete(contract)
    await db.commit()


@router.post("/{contract_id}/submit", response_model=ContractResponse)
async def submit_contract(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """提交合同审查"""
    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    if contract.status != ContractStatus.DRAFT:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="只能提交草稿状态的合同",
        )
    
    contract.status = ContractStatus.PENDING_REVIEW
    await db.commit()
    await db.refresh(contract)
    
    return contract


@router.get("/{contract_id}/versions", response_model=list[ContractVersionResponse])
async def list_contract_versions(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同版本列表"""
    result = await db.execute(
        select(ContractVersion)
        .where(ContractVersion.contract_id == contract_id)
        .order_by(ContractVersion.version_no.desc())
    )
    versions = result.scalars().all()
    return versions


@router.get("/{contract_id}/files", response_model=list[ContractFileResponse])
async def list_contract_files(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同附件列表"""
    result = await db.execute(
        select(ContractFile)
        .where(ContractFile.contract_id == contract_id)
        .order_by(ContractFile.created_at.desc())
    )
    files = result.scalars().all()
    return files


@router.post("/{contract_id}/ai-review")
async def ai_review_contract(
    contract_id: int,
    current_user: User = Depends(require_role("admin", "superadmin", "legal_manager", "legal_specialist")),
    db: AsyncSession = Depends(get_db),
):
    """AI智能审查合同"""
    # 获取合同
    result = await db.execute(select(Contract).where(Contract.id == contract_id))
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    # 提取文件内容（如有）
    file_content = None
    if contract.file_path:
        file_content = await extract_file_content(contract.file_path)
    
    # 构造合同数据
    contract_data = {
        "title": contract.title,
        "contract_type": contract.contract_type.value if contract.contract_type else "",
        "party_a": contract.party_a,
        "party_b": contract.party_b,
        "amount": str(contract.amount) if contract.amount else None,
        "currency": contract.currency,
        "sign_date": contract.sign_date.strftime("%Y-%m-%d") if contract.sign_date else None,
        "effective_date": contract.effective_date.strftime("%Y-%m-%d") if contract.effective_date else None,
        "expiry_date": contract.expiry_date.strftime("%Y-%m-%d") if contract.expiry_date else None,
        "description": contract.description,
        "key_terms": contract.key_terms,
        "special_terms": contract.special_terms,
        "file_path": contract.file_path,
    }
    
    # 调用AI审查
    ai_result = await review_contract_with_ai(contract_data, file_content)
    
    # 更新合同风险信息
    contract.risk_level = ai_result["risk_level"]
    contract.risk_score = ai_result["risk_score"]
    contract.risk_summary = ai_result["summary"]
    
    # 创建审查任务
    review_task = ReviewTask(
        contract_id=contract.id,
        reviewer_id=current_user.id,
        assigned_by=current_user.id,
        status=ReviewTaskStatus.COMPLETED,
        risk_level=ai_result["risk_level"],
        risk_score=ai_result["risk_score"],
        summary=ai_result["summary"],
        review_opinion=ai_result["summary"],
        started_at=datetime.utcnow(),
        completed_at=datetime.utcnow(),
    )
    db.add(review_task)
    await db.flush()
    
    # 创建审查意见
    for finding in ai_result.get("findings", []):
        opinion = ReviewOpinion(
            review_task_id=review_task.id,
            reviewer_id=current_user.id,
            opinion_type=finding.get("category", "risk"),
            content=f"**{finding.get('title', '')}**\n\n{finding.get('description', '')}",
            suggestion=finding.get("suggestion"),
            risk_level=finding.get("risk_level"),
            clause_reference=finding.get("clause_reference"),
            legal_basis=finding.get("legal_basis"),
        )
        db.add(opinion)
    
    # 更新合同状态
    contract.status = ContractStatus.REVIEWED
    contract.reviewed_at = datetime.utcnow()
    contract.reviewer_id = current_user.id
    
    await db.commit()
    await db.refresh(review_task)
    
    return {
        "message": "AI审查完成",
        "review_task_id": review_task.id,
        "risk_level": ai_result["risk_level"],
        "risk_score": ai_result["risk_score"],
        "summary": ai_result["summary"],
        "findings_count": len(ai_result.get("findings", [])),
    }


@router.post("/{contract_id}/modification-suggestions")
async def get_modification_suggestions(
    contract_id: int,
    review_task_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同修改建议"""
    from app.services.contract_modifier import contract_modifier
    
    # 验证合同存在
    contract = await db.get(Contract, contract_id)
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    try:
        suggestions = await contract_modifier.generate_modification_suggestions(
            db, contract_id, review_task_id
        )
        return {
            "contract_id": contract_id,
            "suggestions": [asdict(s) for s in suggestions],
            "total": len(suggestions)
        }
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        import traceback
        print(f"生成修改建议错误: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"生成修改建议失败: {str(e) or '未知错误'}"
        )


@router.post("/{contract_id}/apply-modifications")
async def apply_modifications(
    contract_id: int,
    suggestion_ids: List[str],
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """应用修改建议"""
    from app.services.contract_modifier import contract_modifier
    
    # 验证合同存在
    contract = await db.get(Contract, contract_id)
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    try:
        result = await contract_modifier.apply_modification(
            db, contract_id, suggestion_ids, current_user.id
        )
        return {
            "message": f"已应用 {result.applied_count} 个修改建议",
            "applied_count": result.applied_count,
            "total_suggestions": result.total_suggestions,
            "version_id": result.version_id,
            "modified_content": result.modified_content,
            "diff_summary": result.diff_summary
        }
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"应用修改建议失败: {str(e)}"
        )


@router.get("/{contract_id}/export-modified")
async def export_modified_contract(
    contract_id: int,
    format: str = Query("word", regex="^(word|pdf|markdown)$"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """导出修改后的合同"""
    from app.services.contract_modifier import contract_modifier
    from fastapi.responses import StreamingResponse
    import io
    
    # 获取合同
    contract = await db.get(Contract, contract_id)
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    # 获取最新版本的修改后内容
    result = await db.execute(
        select(ContractVersion)
        .where(ContractVersion.contract_id == contract_id)
        .order_by(ContractVersion.version_no.desc())
        .limit(1)
    )
    latest_version = result.scalar_one_or_none()
    
    # 获取审查发现来生成修改后内容
    review_result = await db.execute(
        select(ReviewTask)
        .where(ReviewTask.contract_id == contract_id)
        .order_by(ReviewTask.created_at.desc())
        .limit(1)
    )
    review_task = review_result.scalar_one_or_none()
    
    # 生成合同内容
    if review_task:
        # 获取审查发现
        findings_result = await db.execute(
            select(ReviewOpinion)
            .where(ReviewOpinion.review_task_id == review_task.id)
        )
        findings = findings_result.scalars().all()
        
        # 生成修改建议
        findings_list = [{
            "id": f.id,
            "clause": f.clause_reference or "",
            "content": f.content,
            "risk_level": f.risk_level or "medium",
            "category": f.opinion_type or "",
            "suggestion": f.suggestion or ""
        } for f in findings]
        suggestions = contract_modifier._generate_with_rules(contract, findings_list)
        
        # 生成修改后合同内容
        modified_content = await contract_modifier.rewrite_contract_with_ai(
            contract, suggestions
        )
    else:
        # 没有审查记录，生成基础合同
        modified_content = contract_modifier._rewrite_with_rules(contract, [])
    
    if format == "markdown":
        return StreamingResponse(
            io.BytesIO(modified_content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=modified_{contract_id}.md"}
        )
    elif format == "word":
        # 生成Word文档
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(11)
        
        # 解析Markdown并添加到Word
        for line in modified_content.split("\n"):
            if line.startswith("# "):
                heading = doc.add_heading(line[2:], level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("> "):
                p = doc.add_paragraph()
                run = p.add_run(line[2:])
                run.italic = True
            elif line.strip():
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=modified_{contract_id}.docx"}
        )
    else:  # pdf
        # 生成PDF文档
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)
        
        # 注册中文字体
        font_paths = [
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",  # Linux
        ]
        
        font_name = "Helvetica"
        for fp in font_paths:
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                font_name = "ChineseFont"
                break
            except Exception:
                continue
        
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle("Chinese", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=14)
        heading_style = ParagraphStyle("ChineseHeading", parent=styles["Heading1"], fontName=font_name, fontSize=14, leading=18)
        
        elements = []
        for line in modified_content.split("\n"):
            if line.startswith("# "):
                elements.append(Paragraph(line[2:], heading_style))
            elif line.startswith("## "):
                elements.append(Paragraph(line[3:], heading_style))
            elif line.startswith("### "):
                elements.append(Paragraph(line[4:], heading_style))
            elif line.strip():
                # 处理Markdown格式
                import re
                line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line)
                line = re.sub(r'\[已修改\]', '<font color="green"><b>[已修改]</b></font>', line)
                elements.append(Paragraph(line, normal_style))
            elements.append(Spacer(1, 0.2*cm))
        
        doc.build(elements)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=modified_{contract_id}.pdf"}
        )


@router.get("/{contract_id}/versions")
async def get_contract_versions(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取合同版本历史"""
    # 验证合同存在
    contract = await db.get(Contract, contract_id)
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    result = await db.execute(
        select(ContractVersion)
        .where(ContractVersion.contract_id == contract_id)
        .order_by(ContractVersion.version_no.desc())
    )
    versions = result.scalars().all()
    
    return {
        "contract_id": contract_id,
        "versions": [
            {
                "id": v.id,
                "version_no": v.version_no,
                "change_summary": v.change_summary,
                "created_at": v.created_at.isoformat() if v.created_at else None
            }
            for v in versions
        ]
    }


@router.get("/{contract_id}/versions/compare")
async def compare_versions(
    contract_id: int,
    version1_id: int = Query(..., description="版本1 ID"),
    version2_id: int = Query(..., description="版本2 ID"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """对比两个版本的差异"""
    from app.services.contract_modifier import version_comparer
    
    try:
        result = await version_comparer.compare_versions(
            db, contract_id, version1_id, version2_id
        )
        return result
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"版本对比失败: {str(e)}"
        )


@router.get("/{contract_id}/compare-original")
async def compare_with_original(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """对比原合同和修改后合同"""
    from app.services.contract_modifier import contract_modifier
    from app.services.file_parser import extract_text_from_file
    
    # 获取合同
    contract = await db.get(Contract, contract_id)
    if not contract:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="合同不存在",
        )
    
    # 获取原合同内容 - 从文件中提取
    original_content = None
    if contract.file_path:
        try:
            original_content = await extract_text_from_file(contract.file_path)
        except Exception as e:
            print(f"提取原合同文件内容失败: {e}")
    
    # 如果无法从文件提取，使用AI生成基础合同内容
    if not original_content:
        # 使用规则引擎生成基础合同（不包含修改建议）
        original_content = contract_modifier._rewrite_with_rules(contract, [])
    
    # 获取最新版本的修改后内容
    result = await db.execute(
        select(ContractVersion)
        .where(ContractVersion.contract_id == contract_id)
        .order_by(ContractVersion.version_no.desc())
        .limit(1)
    )
    latest_version = result.scalar_one_or_none()
    
    modified_content = None
    if latest_version:
        # 获取审查发现来重新生成修改后内容
        review_result = await db.execute(
            select(ReviewTask)
            .where(ReviewTask.contract_id == contract_id)
            .order_by(ReviewTask.created_at.desc())
            .limit(1)
        )
        review_task = review_result.scalar_one_or_none()
        
        if review_task:
            findings_result = await db.execute(
                select(ReviewOpinion)
                .where(ReviewOpinion.review_task_id == review_task.id)
            )
            findings = findings_result.scalars().all()
            
            findings_list = [{
                "id": f.id,
                "clause": f.clause_reference or "",
                "content": f.content,
                "risk_level": f.risk_level or "medium",
                "category": f.opinion_type or "",
                "suggestion": f.suggestion or ""
            } for f in findings]
            suggestions = contract_modifier._generate_with_rules(contract, findings_list)
            modified_content = await contract_modifier.rewrite_contract_with_ai(
                contract, suggestions
            )
    
    return {
        "contract_id": contract_id,
        "contract_title": contract.title,
        "original_content": original_content,
        "modified_content": modified_content,
        "has_modifications": modified_content is not None,
        "version_no": latest_version.version_no if latest_version else None
    }


@router.post("/batch-review")
async def batch_review_contracts(
    contract_ids: List[int],
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """批量审查合同"""
    # 验证合同存在
    for contract_id in contract_ids:
        contract = await db.get(Contract, contract_id)
        if not contract:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"合同 {contract_id} 不存在",
            )
    
    # 启动后台任务
    background_tasks.add_task(
        _batch_review_task,
        contract_ids,
        current_user.id
    )
    
    return {
        "message": f"已启动批量审查任务，共 {len(contract_ids)} 个合同",
        "contract_ids": contract_ids,
        "status": "processing"
    }


async def _batch_review_task(contract_ids: List[int], user_id: int):
    """批量审查后台任务"""
    from app.core.database import AsyncSessionLocal
    
    async with AsyncSessionLocal() as db:
        for contract_id in contract_ids:
            try:
                # 获取合同
                contract = await db.get(Contract, contract_id)
                if not contract:
                    continue
                
                # 构建合同数据
                contract_data = {
                    "title": contract.title,
                    "contract_type": contract.contract_type.value if contract.contract_type else "other",
                    "party_a": contract.party_a,
                    "party_b": contract.party_b,
                    "amount": float(contract.amount) if contract.amount else None,
                    "currency": contract.currency,
                    "description": contract.description,
                    "key_terms": contract.key_terms,
                    "special_terms": contract.special_terms,
                }
                
                # 提取文件内容
                file_content = None
                if contract.file_path:
                    file_content = extract_file_content(contract.file_path)
                
                # AI审查
                ai_result = await review_contract_with_ai(contract_data, file_content)
                
                # 更新合同
                contract.risk_level = ai_result["risk_level"]
                contract.risk_score = ai_result["risk_score"]
                contract.risk_summary = ai_result["summary"]
                
                # 创建审查任务
                review_task = ReviewTask(
                    contract_id=contract.id,
                    reviewer_id=user_id,
                    assigned_by=user_id,
                    status=ReviewTaskStatus.COMPLETED,
                    risk_level=ai_result["risk_level"],
                    risk_score=ai_result["risk_score"],
                    summary=ai_result["summary"],
                    review_opinion=ai_result["summary"],
                    started_at=datetime.utcnow(),
                    completed_at=datetime.utcnow(),
                )
                db.add(review_task)
                await db.flush()
                
                # 创建审查意见
                for finding in ai_result.get("findings", []):
                    opinion = ReviewOpinion(
                        review_task_id=review_task.id,
                        reviewer_id=user_id,
                        opinion_type=finding.get("category", "risk"),
                        content=f"**{finding.get('title', '')}**\n\n{finding.get('description', '')}",
                        suggestion=finding.get("suggestion"),
                        risk_level=finding.get("risk_level"),
                        clause_reference=finding.get("clause_reference"),
                        legal_basis=finding.get("legal_basis"),
                    )
                    db.add(opinion)
                
                # 更新状态
                contract.status = ContractStatus.REVIEWED
                contract.reviewed_at = datetime.utcnow()
                contract.reviewer_id = user_id
                
                await db.commit()
                
            except Exception as e:
                print(f"批量审查合同 {contract_id} 失败: {e}")
                continue
