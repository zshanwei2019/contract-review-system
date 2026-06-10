#!/usr/bin/env python3
"""生成系统介绍Word文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return heading

def add_table_with_header(doc, headers, rows, col_widths=None):
    """添加带表头的表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1A56DB')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if i % 2 == 0:
                set_cell_shading(cell, 'F0F4FF')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def generate_intro_doc():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('贵州西工集团')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('智能合同审查系统')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('系统介绍与技术架构')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(4):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('版本: V2.0\n日期: 2026年6月')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ===== 目录页 =====
    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '一、系统概述',
        '二、核心功能',
        '    2.1 AI智能审查',
        '    2.2 风控规则引擎',
        '    2.3 智能文件处理',
        '    2.4 知识库与自学习',
        '    2.5 合规追踪与报告',
        '三、技术架构',
        '四、技术栈',
        '五、用户角色',
        '六、部署方式',
        '七、核心优势',
        '八、代码规模',
        '九、快速开始',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===== 一、系统概述 =====
    add_heading_styled(doc, '一、系统概述', level=1)

    p = doc.add_paragraph()
    run = p.add_run('智能合同审查系统')
    run.font.bold = True
    p.add_run('是一款基于AI技术的企业级合同风险管理平台，专为贵州西工集团量身打造。系统通过5-Agent协作架构、23条行业风控规则、14种毒丸条款检测等核心技术，实现合同全生命周期的智能化管理。')

    p = doc.add_paragraph()
    p.add_run('系统采用前后端分离架构，前端基于Vue 3 + TypeScript + Element Plus，后端基于FastAPI + Python + PostgreSQL，支持Docker容器化部署，可快速扩展至企业级应用。')

    doc.add_paragraph()

    # 系统亮点表格
    add_heading_styled(doc, '系统亮点', level=2)
    highlights = [
        ['5-Agent协作', '法律/财务/商务/风控/知识五维审查'],
        ['23条行业规则', '覆盖采购/销售/外协/物流/租赁'],
        ['14种毒丸检测', '结构隐藏/语言红旗/行为模式'],
        ['12格式解析', 'PDF/Word/Excel/PPT/图片OCR'],
        ['自学习循环', 'FP-Growth关联规则挖掘'],
    ]
    add_table_with_header(doc, ['特性', '说明'], highlights, [4, 10])

    doc.add_page_break()

    # ===== 二、核心功能 =====
    add_heading_styled(doc, '二、核心功能', level=1)

    # 2.1 AI智能审查
    add_heading_styled(doc, '2.1 AI智能审查', level=2)
    p = doc.add_paragraph()
    p.add_run('系统采用').font.size = Pt(11)
    run = p.add_run('5-Agent协作架构')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p.add_run('，从法律、财务、商务、风控、知识五个维度对合同进行全面审查。')

    doc.add_paragraph()
    agent_data = [
        ['法律合规Agent ⚖️', '30%', '法律风险识别、合规性检查'],
        ['财务风险Agent 💰', '20%', '财务条款风险、付款条件分析'],
        ['商务条件Agent 📋', '15%', '商务条款合理性、交付条件'],
        ['风控规则Agent 🛡️', '20%', '23条行业规则+14种毒丸检测'],
        ['知识图谱Agent 🧠', '15%', '历史案例匹配、相似合同检索'],
    ]
    add_table_with_header(doc, ['Agent名称', '权重', '职责'], agent_data, [5, 2, 7])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('四维加权风险评估模型:')
    run.font.bold = True

    dim_data = [
        ['严重度 (Severity)', '40%', '风险事件造成的损失程度'],
        ['可能性 (Probability)', '25%', '风险事件发生的概率'],
        ['财务敞口 (Financial Exposure)', '20%', '潜在财务损失金额'],
        ['责任不对等 (Responsibility Imbalance)', '15%', '甲乙方责任分配公平性'],
    ]
    add_table_with_header(doc, ['维度', '权重', '说明'], dim_data, [6, 2, 6])

    doc.add_page_break()

    # 2.2 风控规则引擎
    add_heading_styled(doc, '2.2 风控规则引擎', level=2)
    p = doc.add_paragraph()
    p.add_run('系统内置').font.size = Pt(11)
    run = p.add_run('23条行业风控规则 (IR-001~IR-023)')
    run.font.bold = True
    p.add_run('和')
    run = p.add_run('14种毒丸条款检测模式')
    run.font.bold = True
    p.add_run('，覆盖制造业合同的各个风险维度。')

    doc.add_paragraph()
    rules_data = [
        ['采购类', 'IR-001~IR-005', '供应商资质、价格合理性、质量标准、交付风险、付款条件'],
        ['销售类', 'IR-006~IR-010', '客户信用、付款保障、产品责任、售后服务、违约赔偿'],
        ['外协类', 'IR-011~IR-015', '转包风险、技术能力、保密义务、知识产权、验收标准'],
        ['物流类', 'IR-016~IR-019', '运输责任、仓储风险、货物保险、延迟交付'],
        ['租赁类', 'IR-020~IR-023', '租金支付、维修责任、转租限制、合同解除'],
    ]
    add_table_with_header(doc, ['类别', '规则编号', '覆盖范围'], rules_data, [3, 3, 8])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('毒丸条款检测:')
    run.font.bold = True

    pp_data = [
        ['结构隐藏型', 'PP-S1~S5', '违约陷阱、赔偿上限、单方终止、自动续约、竞业禁止'],
        ['语言红旗型', 'PP-L1~L8', '模糊表述、单方权利、无限责任、霸王条款、连坐条款、期限陷阱、管辖陷阱、排除救济'],
        ['行为模式型', 'PP-B1', '阴阳合同、异常付款路径'],
    ]
    add_table_with_header(doc, ['类型', '编号', '检测内容'], pp_data, [3, 3, 8])

    doc.add_page_break()

    # 2.3 智能文件处理
    add_heading_styled(doc, '2.3 智能文件处理', level=2)
    p = doc.add_paragraph()
    p.add_run('系统支持').font.size = Pt(11)
    run = p.add_run('12种文件格式')
    run.font.bold = True
    p.add_run('解析，并集成OCR识别能力，可处理扫描件和图片。')

    doc.add_paragraph()
    file_data = [
        ['PDF', 'PyMuPDF', '文本提取+OCR'],
        ['Word', 'python-docx', '全文解析'],
        ['Excel', 'openpyxl', '表格数据提取'],
        ['PowerPoint', 'python-pptx', '幻灯片内容'],
        ['图片', 'Tesseract OCR', '中英文识别(chi_sim+eng)'],
        ['其他', '内置解析', 'CSV/TXT/RTF/HTML/Markdown/XMind'],
    ]
    add_table_with_header(doc, ['格式', '解析引擎', '能力'], file_data, [3, 4, 5])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('智能分类与要素提取:')
    run.font.bold = True

    p = doc.add_paragraph('• 合同自动分类: 6大类17子类，关键词+正则匹配，置信度评分')
    p = doc.add_paragraph('• 关键要素提取: 13项核心字段，正则+LLM双模式')
    p = doc.add_paragraph('• 条款分割引擎: 第X条/数字序号/段落三级分割，逐条风险分析')

    doc.add_page_break()

    # 2.4 知识库与自学习
    add_heading_styled(doc, '2.4 知识库与自学习', level=2)

    p = doc.add_paragraph()
    run = p.add_run('企业知识沉淀:')
    run.font.bold = True
    doc.add_paragraph('• 审查案例自动归档，形成企业审查知识库')
    doc.add_paragraph('• 风险模式自动提取，高频风险规则化')
    doc.add_paragraph('• 人工反馈闭环，持续优化审查策略')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('FP-Growth关联规则挖掘:')
    run.font.bold = True
    doc.add_paragraph('• 基于历史审查数据，发现风险项之间的隐藏关联')
    doc.add_paragraph('• 自动生成关联规则，辅助风险预判')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('FAISS向量索引:')
    run.font.bold = True
    doc.add_paragraph('• 合同文本向量化，支持语义相似度检索')
    doc.add_paragraph('• 相似案例智能匹配，辅助审查决策')

    doc.add_page_break()

    # 2.5 合规追踪与报告
    add_heading_styled(doc, '2.5 合规追踪与报告', level=2)

    p = doc.add_paragraph()
    run = p.add_run('合规检查清单:')
    run.font.bold = True
    doc.add_paragraph('• 6类合同合规检查模板（采购/销售/外协/保密/服务/租赁）')
    doc.add_paragraph('• 合规率自动评估，合规缺口识别')
    doc.add_paragraph('• 整改计划自动生成')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('报告导出:')
    run.font.bold = True
    doc.add_paragraph('• Word格式: 完整审查报告，含表格和格式化')
    doc.add_paragraph('• PDF格式: 正式报告，适合归档和打印')
    doc.add_paragraph('• 降级方案: 纯文本报告（无依赖环境）')

    doc.add_page_break()

    # ===== 三、技术架构 =====
    add_heading_styled(doc, '三、技术架构', level=1)

    arch_desc = """系统采用前后端分离的微服务架构，分为以下五层:

┌─────────────────────────────────────────────┐
│           前端层 (Frontend)                    │
│     Vue 3 + TypeScript + Element Plus         │
└─────────────────────────────────────────────┘
                    │
                    ▼
┌─────────────────────────────────────────────┐
│           API网关层 (Gateway)                  │
│        Nginx 反向代理 + 负载均衡               │
└─────────────────────────────────────────────┘
                    │
                    ▼
┌─────────────────────────────────────────────┐
│           后端服务层 (Backend)                 │
│    FastAPI + Python 3.14 + SQLAlchemy 2.0     │
│    ┌─────────────────────────────────────┐    │
│    │  API模块: auth/users/contracts/     │    │
│    │          reviews/risks/workflows/    │    │
│    │          agent (AI智能体核心)        │    │
│    └─────────────────────────────────────┘    │
└─────────────────────────────────────────────┘
                    │
                    ▼
┌─────────────────────────────────────────────┐
│           AI服务层 (AI Services)               │
│    ┌─────────────────────────────────────┐    │
│    │  • multi_agent.py    5-Agent引擎    │    │
│    │  • risk_rules_engine 风控规则引擎   │    │
│    │  • ai_review.py      AI审查服务     │    │
│    │  • vector_index.py   FAISS索引      │    │
│    │  • self_learning.py  自学习循环     │    │
│    └─────────────────────────────────────┘    │
└─────────────────────────────────────────────┘
                    │
                    ▼
┌─────────────────────────────────────────────┐
│           数据层 (Data Layer)                  │
│    PostgreSQL 15 + Redis 7 + FAISS + 文件系统 │
└─────────────────────────────────────────────┘"""

    p = doc.add_paragraph()
    run = p.add_run(arch_desc)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===== 四、技术栈 =====
    add_heading_styled(doc, '四、技术栈', level=1)

    tech_data = [
        ['前端', 'Vue 3', '响应式UI框架'],
        ['前端', 'TypeScript', '类型安全'],
        ['前端', 'Element Plus', '企业级组件库'],
        ['前端', 'Vite', '构建工具'],
        ['前端', 'Pinia', '状态管理'],
        ['前端', 'ECharts', '数据可视化'],
        ['后端', 'FastAPI', '高性能异步框架'],
        ['后端', 'Python 3.14', '最新语言特性'],
        ['后端', 'SQLAlchemy 2.0', '异步ORM'],
        ['后端', 'Pydantic 2.10', '数据验证'],
        ['后端', 'JWT + bcrypt', '安全认证'],
        ['数据库', 'PostgreSQL 15', '主数据库'],
        ['数据库', 'Redis 7', '缓存+会话'],
        ['数据库', 'FAISS', '向量检索'],
        ['AI', 'OpenAI API', 'GPT-4兼容接口'],
        ['AI', 'sentence-transformers', '文本向量化'],
        ['AI', 'Tesseract OCR', '图片文字识别'],
        ['部署', 'Docker', '容器化部署'],
        ['部署', 'Nginx', '反向代理'],
        ['部署', 'Alembic', '数据库迁移'],
    ]
    add_table_with_header(doc, ['层级', '技术', '说明'], tech_data, [3, 5, 6])

    doc.add_page_break()

    # ===== 五、用户角色 =====
    add_heading_styled(doc, '五、用户角色', level=1)

    p = doc.add_paragraph()
    p.add_run('系统支持').font.size = Pt(11)
    run = p.add_run('7种用户角色')
    run.font.bold = True
    p.add_run('，实现细粒度的权限控制。')

    doc.add_paragraph()
    role_data = [
        ['超级管理员', '系统全部权限，包括系统配置和数据管理'],
        ['管理员', '用户管理、系统配置、数据备份'],
        ['法务经理', '审查任务分配、结果审核、规则管理'],
        ['法务专员', '执行审查、提交意见、案例归档'],
        ['业务经理', '合同创建、流程发起、进度跟踪'],
        ['业务专员', '合同录入、进度查看、资料上传'],
        ['高管层', '数据看板、决策支持、报表查看'],
    ]
    add_table_with_header(doc, ['角色', '权限说明'], role_data, [4, 10])

    doc.add_page_break()

    # ===== 六、部署方式 =====
    add_heading_styled(doc, '六、部署方式', level=1)

    add_heading_styled(doc, '6.1 开发环境', level=2)
    code1 = """# 后端启动
cd backend
pip install -r requirements.txt
python -m app.main

# 前端启动
cd frontend
npm install
npm run dev"""
    p = doc.add_paragraph()
    run = p.add_run(code1)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    add_heading_styled(doc, '6.2 生产环境 (Docker Compose)', level=2)
    code2 = """# 一键部署
docker-compose -f docker-compose.server.yml up -d

# 服务地址
前端: http://localhost:3007
后端API: http://localhost:8000
API文档: http://localhost:8000/docs"""
    p = doc.add_paragraph()
    run = p.add_run(code2)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # ===== 七、核心优势 =====
    add_heading_styled(doc, '七、核心优势', level=1)

    advantages = [
        ['专业垂直', '深度适配制造业合同场景，23条行业专属风控规则，14种毒丸条款精准识别'],
        ['智能协作', '5-Agent五维加权审查，多步推理链自主决策，知识图谱辅助判断'],
        ['持续进化', '人工反馈闭环学习，FP-Growth关联规则挖掘，风险模式自动沉淀'],
        ['企业级可靠', '7角色细粒度权限，JWT安全认证，审计日志追踪，Docker容器化部署'],
        ['降级容灾', '无AI环境自动降级为规则审查，确保系统可用性'],
    ]
    add_table_with_header(doc, ['优势', '说明'], advantages, [4, 10])

    doc.add_page_break()

    # ===== 八、代码规模 =====
    add_heading_styled(doc, '八、代码规模', level=1)

    scale_data = [
        ['后端服务', '15+', '8000+'],
        ['前端页面', '20+', '6000+'],
        ['API接口', '30+', '-'],
        ['风控规则', '37条', '-'],
        ['总计', '50+', '15000+'],
    ]
    add_table_with_header(doc, ['模块', '文件数', '代码行数'], scale_data, [4, 3, 3])

    doc.add_page_break()

    # ===== 九、快速开始 =====
    add_heading_styled(doc, '九、快速开始', level=1)

    steps = [
        ('1. 克隆代码', 'git clone https://github.com/zshanwei2019/contract-review-system.git'),
        ('2. 启动后端', 'cd backend\npip install -r requirements.txt\npython -m app.main'),
        ('3. 启动前端', 'cd frontend\nnpm install\nnpm run dev'),
        ('4. 访问系统', '地址: http://localhost:5175\n账号: admin / admin123'),
    ]

    for title, code in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        doc.add_paragraph()

    # 保存
    output_path = os.path.join(os.path.dirname(__file__), '贵州西工集团智能合同审查系统-介绍文档.docx')
    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    generate_intro_doc()
